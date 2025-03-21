import os
import io
import requests
import fitz  # PyMuPDF
import docx
from docx import Document
from docx.shared import RGBColor
from flask import Flask, render_template, request, send_file, redirect, url_for
from werkzeug.utils import secure_filename
from docx.shared import Pt
from dotenv import load_dotenv

# Carrega variáveis de ambiente do .env
load_dotenv()

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
ALLOWED_EXTENSIONS = {'pdf', 'docx'}

# Obtenha as chaves de API do ambiente (.env)
DEEPSEEK_API_KEY = os.getenv('DEEPSEEK_API_KEY')
SERPAPI_API_KEY = os.getenv('SERPAPI_API_KEY')

# Variável global (exemplo didático) para armazenar o último DOCX gerado
DOC_BUFFER = None

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(pdf_path):
    text = ""
    with fitz.open(pdf_path) as doc:
        for page in doc:
            text += page.get_text()
    return text

def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs])

def get_source_from_serpapi(query):
    url = "https://serpapi.com/search"
    params = {
        "engine": "google",
        "q": query,
        "api_key": SERPAPI_API_KEY,
        "num": 1
    }
    try:
        response = requests.get(url, params=params)
        response.raise_for_status()
        data = response.json()
        if "organic_results" in data and len(data["organic_results"]) > 0:
            return data["organic_results"][0].get("link", "Fonte não encontrada")
    except requests.RequestException:
        pass
    return "Fonte não encontrada"

def analyze_plagiarism_with_source(text):
    """
    Retorna (resultado_bruto, trechos_plagio)
    onde:
      - resultado_bruto: string formatada para exibição
      - trechos_plagio: lista de dicionários { "trecho": ..., "fonte": ... }
    """
    url = "https://api.deepseek.com/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
        "Content-Type": "application/json"
    }
    data = {
        "model": "deepseek-chat",
        "messages": [
            {"role": "system", "content": "Você é um assistente que detecta plágio comparando textos."},
            {"role": "user", "content": (
                "Verifique se esse texto possui plágio na internet e destaque os trechos plagiados. "
                "Para cada trecho, informe também a fonte original (por exemplo, a URL ou outra referência). "
                "\n\n" + text
            )}
        ]
    }

    response = requests.post(url, headers=headers, json=data)
    if response.status_code != 200:
        return f"Erro na API DeepSeek (status: {response.status_code})", []

    deepseek_result = response.json()["choices"][0]["message"]["content"]
    resultado_bruto = ""
    trechos_plagio = []

    linhas = deepseek_result.splitlines()
    for linha in linhas:
        if linha.strip().startswith("Trecho"):
            # Exemplo de formato: "Trecho: ABC - Fonte: https://..."
            try:
                after_trecho = linha.split("Trecho:", 1)[1].strip()
                trecho, fonte_part = after_trecho.split("- Fonte:", 1)
                trecho = trecho.strip()
                fonte_part = fonte_part.strip()

                # Se quiser buscar a fonte exata no SerpApi, descomente:
                # fonte = get_source_from_serpapi(trecho)
                fonte = fonte_part

                trechos_plagio.append({
                    "trecho": trecho,
                    "fonte": fonte
                })

                resultado_bruto += f"Trecho: {trecho} - Fonte: {fonte}\n"
            except:
                resultado_bruto += linha + "\n"
        else:
            resultado_bruto += linha + "\n"

    return resultado_bruto, trechos_plagio

def highlight_plagiarized_in_docx(original_text, trechos_plagio):
    """
    Cria um Document e destaca os trechos plagiados em vermelho.
    Retorna o objeto Document.
    """
    doc = Document()
    # Ordena os trechos do maior para o menor (para evitar substituições parciais)
    trechos_plagio_ordenados = sorted(trechos_plagio, key=lambda x: len(x["trecho"]), reverse=True)

    p = doc.add_paragraph()
    normal_run = p.add_run("")
    current_text = original_text

    for t in trechos_plagio_ordenados:
        trecho = t["trecho"]
        if trecho in current_text:
            before, match, after = current_text.partition(trecho)
            normal_run.add_text(before)
            # Cria um run para o trecho plagiado
            plagio_run = p.add_run(trecho)
            font = plagio_run.font
            font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # vermelho
            font.bold = True
            current_text = after
    # Adiciona o resto do texto
    normal_run.add_text(current_text)

    return doc

def calculate_plagiarism_percentage(original_text, trechos_plagio):
    """
    Calcula um percentual de plágio simples,
    contando palavras plagiadas em relação ao total de palavras.
    """
    total_words = len(original_text.split())
    if total_words == 0:
        return 0.0

    plag_words = 0
    for item in trechos_plagio:
        plag_words += len(item["trecho"].split())

    return round((plag_words / total_words) * 100, 2)

@app.route("/")
def home():
    return render_template("index.html")

@app.route("/comparacao_texto", methods=["GET", "POST"])
def comparacao_texto():
    text1, text2, ai_analysis = "", "", ""
    if request.method == "POST":
        text1 = request.form.get("text1", "")
        text2 = request.form.get("text2", "")
        comparison_text = f"Comparação:\nTexto 1:\n{text1}\n\nTexto 2:\n{text2}"
        resultado_bruto, trechos_plagio = analyze_plagiarism_with_source(comparison_text)
        ai_analysis = resultado_bruto
    return render_template("comparacao.html", text1=text1, text2=text2, ai_analysis=ai_analysis)

@app.route("/verificacao_plagio", methods=["GET", "POST"])
def verificacao_plagio():
    global DOC_BUFFER
    text, plagio_result = "", ""
    percentage = 0.0
    download_ready = False

    if request.method == "POST":
        file_text = ""
        if "file" in request.files and request.files["file"].filename != "":
            file = request.files["file"]
            if allowed_file(file.filename):
                filename = secure_filename(file.filename)
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(file_path)
                if filename.lower().endswith(".pdf"):
                    file_text = extract_text_from_pdf(file_path)
                elif filename.lower().endswith(".docx"):
                    file_text = extract_text_from_docx(file_path)

        text = request.form.get("text", "") or file_text

        if text.strip():
            # Faz a análise
            resultado_bruto, trechos_plagio = analyze_plagiarism_with_source(text)
            plagio_result = resultado_bruto

            # Calcula percentual
            percentage = calculate_plagiarism_percentage(text, trechos_plagio)

            # Gera docx com trechos em vermelho
            doc = highlight_plagiarized_in_docx(text, trechos_plagio)
            temp_file = io.BytesIO()
            doc.save(temp_file)
            temp_file.seek(0)

            # Armazena em DOC_BUFFER (exemplo simples para poder baixar)
            DOC_BUFFER = temp_file.read()
            download_ready = True

    return render_template(
        "plagio.html",
        text=text,
        plagio_result=plagio_result,
        percentage=percentage,
        download_ready=download_ready
    )

@app.route("/download_doc")
def download_doc():
    """
    Retorna o último DOCX gerado em /verificacao_plagio,
    armazenado na variável global DOC_BUFFER.
    """
    global DOC_BUFFER
    if not DOC_BUFFER:
        return "Nenhum arquivo gerado ainda", 400

    return send_file(
        io.BytesIO(DOC_BUFFER),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name="resultado_plagio.docx"
    )

if __name__ == "__main__":
    if not os.path.exists(app.config['UPLOAD_FOLDER']):
        os.makedirs(app.config['UPLOAD_FOLDER'])
    app.run(debug=True)
